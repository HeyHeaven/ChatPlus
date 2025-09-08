import io
from datetime import datetime
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.units import cm
from reportlab.lib import colors
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.title_style = ParagraphStyle('CustomTitle', parent=self.styles['Heading1'], fontSize=22, spaceAfter=24, alignment=TA_CENTER, textColor=colors.darkblue)
        self.section_style = ParagraphStyle('CustomSection', parent=self.styles['Heading2'], fontSize=14, spaceAfter=12, textColor=colors.darkgreen)
        self.body_style = ParagraphStyle('CustomBody', parent=self.styles['Normal'], fontSize=11, spaceAfter=10, alignment=TA_LEFT)

    def _metric_table(self, kpis, detected_format):
        data = [
            ['Metric','Value'],
            ['Total Messages', kpis.get('total_messages','0')],
            ['Total Words', kpis.get('total_words','0')],
            ['Media Messages', kpis.get('media_messages','0')],
            ['Links Shared', kpis.get('links_shared','0')],
            ['Export Format', detected_format],
            ['Analysis Period', kpis.get('date_range','N/A')],
        ]
        tbl = Table(data, colWidths=[7*cm, 7*cm])
        tbl.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0),colors.grey),
            ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
            ('GRID',(0,0),(-1,-1),0.5,colors.black),
        ]))
        return tbl

    def _chart_image_from_df(self, fig_builder):
        buf = io.BytesIO()
        fig = fig_builder()
        fig.savefig(buf, format='png', dpi=220, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

    def generate_pdf_report(self, analysis_data, selected_user, detected_format, charts_data):
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=24)
        story = []
        story.append(Paragraph("WhatsApp Chat Analysis Report", self.title_style))
        story.append(Paragraph(f"Generated on: {datetime.now():%B %d, %Y %I:%M %p}", self.body_style))
        story.append(Spacer(1,12))
        story.append(Paragraph("Key Metrics", self.section_style))
        story.append(self._metric_table(analysis_data, detected_format))
        story.append(Spacer(1,12))

        # Timeline
        if charts_data.get("timeline") is not None:
            story.append(Paragraph("Timeline (Monthly)", self.section_style))
            tdf = charts_data["timeline"].copy()
            def _build():
                import matplotlib.pyplot as plt
                fig, ax = plt.subplots(figsize=(7.5,4))
                ax.plot(tdf["time"], tdf["message"], color="green", marker="o", linewidth=2)
                ax.set_xlabel("Month-Year"); ax.set_ylabel("Messages"); ax.grid(True, alpha=0.3)
                plt.xticks(rotation=45); plt.tight_layout(); return fig
            img_buf = self._chart_image_from_df(_build)
            story.append(Image(img_buf, width=16*cm, height=9*cm))
            story.append(Spacer(1,10))

        # Users
        if charts_data.get("user_activity") is not None:
            story.append(Paragraph("User Activity", self.section_style))
            x = charts_data["user_activity"]
            def _build():
                import matplotlib.pyplot as plt
                fig, ax = plt.subplots(figsize=(7.5,4))
                ax.bar(range(len(x)), x.values, color='skyblue', alpha=0.85)
                ax.set_xticks(range(len(x))); ax.set_xticklabels(x.index, rotation=45, ha='right')
                ax.set_xlabel("Users"); ax.set_ylabel("Messages"); ax.grid(True, alpha=0.2); plt.tight_layout(); return fig
            img_buf = self._chart_image_from_df(_build)
            story.append(Image(img_buf, width=16*cm, height=9*cm))
            story.append(Spacer(1,10))

        # Words (HARDENED)
        if charts_data.get("word_analysis") is not None:
            story.append(Paragraph("Word Frequency", self.section_style))
            wdf = charts_data["word_analysis"].copy().reset_index(drop=True)
            if not wdf.empty and len(wdf.columns) >= 2:
                wdf = wdf.iloc[:, :2].head(15)
                wdf.columns = ["word", "count"]
                def _build():
                    import matplotlib.pyplot as plt
                    fig, ax = plt.subplots(figsize=(7.5,5))
                    ax.barh(wdf["word"].astype(str), wdf["count"].astype(float), color="steelblue", alpha=0.8)
                    ax.set_xlabel("Frequency"); ax.set_ylabel("Words")
                    ax.grid(True, alpha=0.2); plt.tight_layout(); return fig
                img_buf = self._chart_image_from_df(_build)
                story.append(Image(img_buf, width=16*cm, height=10*cm))
                story.append(Spacer(1,10))

        # Emojis
        if charts_data.get("emoji_analysis") is not None:
            story.append(Paragraph("Emoji Usage", self.section_style))
            edf = charts_data["emoji_analysis"].head(8).copy()
            def _build():
                import matplotlib.pyplot as plt
                fig, ax = plt.subplots(figsize=(6.5,6))
                ax.pie(edf.iloc[:,1], labels=edf.iloc[:,0], autopct='%1.1f%%', startangle=90)
                ax.set_title("Top Emojis"); plt.tight_layout(); return fig
            img_buf = self._chart_image_from_df(_build)
            story.append(Image(img_buf, width=12*cm, height=12*cm))
            story.append(Spacer(1,10))

        # AI summary
        if analysis_data.get("ai_summary"):
            story.append(Paragraph("AI Summary", self.section_style))
            story.append(Paragraph(analysis_data["ai_summary"].replace("**","").replace("🤖",""), self.body_style))

        doc.build(story)
        buf.seek(0)
        return buf

    def generate_docx_report(self, analysis_data, selected_user, detected_format, charts_data):
        buf = io.BytesIO()
        doc = Document()

        # Title
        title = doc.add_paragraph("WhatsApp Chat Analysis Report")
        title.runs[0].bold = True
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generated date
        date_par = doc.add_paragraph(f"Generated on: {datetime.now():%B %d, %Y %I:%M %p}")
        date_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        # Key Metrics
        doc.add_paragraph("Key Metrics").runs[0].bold = True
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        metrics = [
            ("Total Messages", str(analysis_data.get("total_messages", "0"))),
            ("Total Words", str(analysis_data.get("total_words", "0"))),
            ("Media Messages", str(analysis_data.get("media_messages", "0"))),
            ("Links Shared", str(analysis_data.get("links_shared", "0"))),
            ("Export Format", str(detected_format)),
            ("Analysis Period", str(analysis_data.get("date_range", "N/A"))),
        ]
        for label, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph("")

        # Word Analysis (optional)
        wdf = charts_data.get("word_analysis")
        if wdf is not None and not wdf.empty and len(wdf.columns) >= 2:
            doc.add_paragraph("Top Words").runs[0].bold = True
            sub_table = doc.add_table(rows=1, cols=2)
            sh = sub_table.rows[0].cells
            sh[0].text = "Word"; sh[1].text = "Count"
            subset = wdf.iloc[:, :2].copy()
            subset.columns = ["word", "count"]
            for _, r in subset.head(15).iterrows():
                rc = sub_table.add_row().cells
                rc[0].text = str(r["word"]) 
                rc[1].text = str(int(r["count"])) if str(r["count"]).isdigit() else str(r["count"]) 

            doc.add_paragraph("")

        # AI Summary (optional)
        if analysis_data.get("ai_summary"):
            doc.add_paragraph("AI Summary").runs[0].bold = True
            doc.add_paragraph(str(analysis_data["ai_summary"]).replace("**",""))

        doc.save(buf)
        buf.seek(0)
        return buf
